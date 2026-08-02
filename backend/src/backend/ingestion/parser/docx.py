from importlib import metadata
from pathlib import Path
from typing import Any

from backend.ingestion.parser.errors import ParserDependencyError
from backend.ingestion.parser.models import ParsedBlock, ParsedDocument, ParsedPage


class DocxParser:
    """DOCX parser adapter backed by python-docx."""

    parser_name = "python_docx"
    supported_extensions = frozenset({".docx"})
    supported_mime_types = frozenset(
        {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
    )

    def parse(self, file_path: Path) -> ParsedDocument:
        """Extract paragraphs and table rows from a DOCX document.

        Args:
            file_path: Absolute or relative path to a DOCX file stored by ingestion.

        Returns:
            A ParsedDocument containing one logical page because DOCX files do not
            expose dependable page boundaries without a rendering engine.
        """
        try:
            from docx import Document
        except ImportError as exc:
            raise ParserDependencyError(self.parser_name, "python-docx") from exc

        document = Document(file_path)
        blocks = self._parse_blocks(document)
        document_text = "\n\n".join(block.text for block in blocks)

        # Represent the DOCX as one logical page because pagination depends on fonts
        # and rendering settings that python-docx does not calculate.
        page = ParsedPage(
            page_number=1,
            text=document_text,
            blocks=blocks,
            metadata={"logical_page": True},
        )

        return ParsedDocument(
            text=document_text,
            pages=[page],
            metadata={"source_path": str(file_path)},
            parser_name=self.parser_name,
            parser_version=self._get_parser_version(),
        )

    def _parse_blocks(self, document: Any) -> list[ParsedBlock]:
        """Convert DOCX paragraphs and table rows into ordered text blocks.

        Args:
            document: Open python-docx Document instance to extract from.

        Returns:
            A list of ParsedBlock instances in the document body order.
        """
        paragraphs_by_element = {
            paragraph._element: paragraph for paragraph in document.paragraphs
        }
        tables_by_element = {table._element: table for table in document.tables}
        blocks: list[ParsedBlock] = []
        table_index = 0

        # Iterate body XML elements so paragraphs and tables retain their source order.
        for body_element in document.element.body.iterchildren():
            paragraph = paragraphs_by_element.get(body_element)

            # Convert a body paragraph into a block when it contains readable text.
            if paragraph is not None:
                text = paragraph.text.strip()

                if text:
                    blocks.append(
                        ParsedBlock(
                            text=text,
                            page_number=1,
                            block_index=len(blocks),
                            metadata={"block_type": "paragraph"},
                        )
                    )

                continue

            table = tables_by_element.get(body_element)

            # Convert each non-empty table row into one block for later chunking.
            if table is not None:
                self._append_table_rows(table, table_index, blocks)
                table_index += 1

        return blocks

    def _append_table_rows(
        self,
        table: Any,
        table_index: int,
        blocks: list[ParsedBlock],
    ) -> None:
        """Append readable rows from one DOCX table to the parsed block list.

        Args:
            table: Open python-docx Table instance whose rows should be extracted.
            table_index: Zero-based ordinal of the table in the document body.
            blocks: Mutable parsed-block list that receives table row blocks.

        Returns:
            None. Rows are appended directly to the supplied block list.
        """
        # Emit one normalized block per visible row so table values retain row context.
        for row_index, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("\n", " | ") for cell in row.cells]
            text = " | ".join(cell for cell in cells if cell)

            if text:
                blocks.append(
                    ParsedBlock(
                        text=text,
                        page_number=1,
                        block_index=len(blocks),
                        metadata={
                            "block_type": "table_row",
                            "table_index": table_index,
                            "row_index": row_index,
                        },
                    )
                )

    def _get_parser_version(self) -> str:
        """Read the installed python-docx package version.

        Args:
            None.

        Returns:
            The installed python-docx version, or "unknown" when unavailable.
        """
        try:
            return metadata.version("python-docx")
        except metadata.PackageNotFoundError:
            return "unknown"
